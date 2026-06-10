from __future__ import annotations

import fitz
from docx import Document
from posejdon_docs.parsers.json_parser import JSONParser
from posejdon_docs.parsers.xml_parser import XMLParser

from posejdon.core.enums import DocumentKind
from posejdon.domain.reports import ValidationResult


class StructuralValidator:
    def validate(
        self, *, input_path: str, output_path: str, document_kind: DocumentKind
    ) -> ValidationResult:
        warnings: list[str] = []
        errors: list[str] = []
        checks: list[str] = []

        try:
            if document_kind == DocumentKind.DOCX:
                original = Document(input_path)
                output = Document(output_path)
                checks.append("docx_opened")
                original_body = len(original.paragraphs)
                output_body = len(output.paragraphs)
                if abs(original_body - output_body) > 2:
                    warnings.append("Paragraph count changed beyond tolerance.")
                checks.append("docx_paragraph_count_checked")
                if len(original.tables) != len(output.tables):
                    errors.append("DOCX table count changed unexpectedly.")
                checks.append("docx_table_count_checked")
                original_headers = sum(
                    len(section.header.paragraphs) for section in original.sections
                )
                output_headers = sum(len(section.header.paragraphs) for section in output.sections)
                if original_headers != output_headers:
                    warnings.append("Header paragraph count changed.")
                checks.append("docx_header_count_checked")
                original_footers = sum(
                    len(section.footer.paragraphs) for section in original.sections
                )
                output_footers = sum(len(section.footer.paragraphs) for section in output.sections)
                if original_footers != output_footers:
                    warnings.append("Footer paragraph count changed.")
                checks.append("docx_footer_count_checked")
            elif document_kind == DocumentKind.PDF:
                original_pdf = fitz.open(input_path)
                output_pdf = fitz.open(output_path)
                try:
                    checks.append("pdf_opened")
                    if original_pdf.page_count != output_pdf.page_count:
                        errors.append("Page count changed unexpectedly.")
                    checks.append("pdf_page_count_checked")
                    for page_index in range(original_pdf.page_count):
                        original_page = original_pdf.load_page(page_index)
                        output_page = output_pdf.load_page(page_index)
                        if not output_page.get_text("blocks") and original_page.get_text("blocks"):
                            warnings.append(f"PDF page {page_index} no longer exposes text blocks.")
                    checks.append("pdf_page_block_presence_checked")
                finally:
                    original_pdf.close()
                    output_pdf.close()
            elif document_kind == DocumentKind.JSON:
                original = JSONParser().parse(input_path)
                output = JSONParser().parse(output_path)
                checks.append("json_opened")
                if len(original.text_segments) != len(output.text_segments):
                    warnings.append("JSON scalar segment count changed.")
                checks.append("json_segment_count_checked")
            elif document_kind == DocumentKind.XML:
                original = XMLParser().parse(input_path)
                output = XMLParser().parse(output_path)
                checks.append("xml_opened")
                if len(original.text_segments) != len(output.text_segments):
                    warnings.append("XML text and attribute segment count changed.")
                checks.append("xml_segment_count_checked")
        except Exception as exc:
            errors.append(str(exc))

        return ValidationResult(
            passed=not errors,
            structure_checks=checks,
            warnings=warnings,
            errors=errors,
        )
