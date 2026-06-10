from __future__ import annotations

import re

import fitz
from docx import Document
from posejdon_docs.parsers.base import ParsedTextSegment
from posejdon_docs.parsers.json_parser import JSONParser
from posejdon_docs.parsers.xml_parser import XMLParser

from posejdon.core.enums import DocumentKind
from posejdon.domain.entities import SensitiveEntity
from posejdon.domain.reports import LeakageScanResult, SegmentLeakageFinding


class LeakageValidator:
    _CONNECTOR_CHARS = r"\w@._-"

    def extract_segments(self, path: str, document_kind: DocumentKind) -> list[ParsedTextSegment]:
        if document_kind == DocumentKind.DOCX:
            document = Document(path)
            segments: list[ParsedTextSegment] = []

            def add_paragraphs(paragraphs, prefix: str) -> None:
                for index, paragraph in enumerate(paragraphs):
                    text = paragraph.text
                    if not text:
                        continue
                    segment_id = f"{prefix}:p:{index}"
                    segments.append(
                        ParsedTextSegment(
                            segment_id=segment_id,
                            text=text,
                            container_id=segment_id,
                            section_id=segment_id,
                            start_offset=0,
                            end_offset=len(text),
                        )
                    )

            add_paragraphs(document.paragraphs, "body")
            for table_index, table in enumerate(document.tables):
                for row_index, row in enumerate(table.rows):
                    for cell_index, cell in enumerate(row.cells):
                        add_paragraphs(
                            cell.paragraphs,
                            f"table:{table_index}:r:{row_index}:c:{cell_index}",
                        )
            for section_index, section in enumerate(document.sections):
                add_paragraphs(section.header.paragraphs, f"header:{section_index}")
                add_paragraphs(section.footer.paragraphs, f"footer:{section_index}")
            return segments

        if document_kind == DocumentKind.JSON:
            return JSONParser().parse(path).text_segments

        if document_kind == DocumentKind.XML:
            return XMLParser().parse(path).text_segments

        pdf = fitz.open(path)
        try:
            segments: list[ParsedTextSegment] = []
            for page_index in range(pdf.page_count):
                page = pdf.load_page(page_index)
                for block_index, block in enumerate(page.get_text("blocks")):
                    x0, y0, x1, y1, text, *_ = block
                    clean = text.strip()
                    if not clean:
                        continue
                    segment_id = f"page:{page_index}:block:{block_index}"
                    segments.append(
                        ParsedTextSegment(
                            segment_id=segment_id,
                            text=clean,
                            container_id=segment_id,
                            page_index=page_index,
                            section_id=segment_id,
                            start_offset=0,
                            end_offset=len(clean),
                        )
                    )
            return segments
        finally:
            pdf.close()

    def extract_text(self, path: str, document_kind: DocumentKind) -> str:
        return "\n".join(
            segment.text for segment in self.extract_segments(path, document_kind)
        ).strip()

    def validate(
        self,
        *,
        output_path: str,
        document_kind: DocumentKind,
        entities: list[SensitiveEntity],
    ) -> LeakageScanResult:
        output_segments = self.extract_segments(output_path, document_kind)
        segment_lookup = {segment.segment_id: segment.text for segment in output_segments}
        global_findings: set[str] = set()
        normalized_findings: set[str] = set()
        segmented_findings: list[SegmentLeakageFinding] = []

        for segment in output_segments:
            matched = sorted(
                {
                    entity.raw_text
                    for entity in entities
                    if self._contains_surface(segment.text, entity.raw_text)
                    and (entity.segment_id is None or entity.segment_id == segment.segment_id)
                }
            )
            normalized_matched = sorted(
                {
                    entity.raw_text
                    for entity in entities
                    if self._normalized_contains(segment.text, entity.raw_text)
                    and (entity.segment_id is None or entity.segment_id == segment.segment_id)
                }
            )
            if matched:
                segmented_findings.append(
                    SegmentLeakageFinding(segment_id=segment.segment_id, findings=matched)
                )
                global_findings.update(matched)
            if normalized_matched:
                normalized_findings.update(normalized_matched)

        full_text = "\n".join(segment_lookup.values())
        for entity in entities:
            if self._contains_surface(full_text, entity.raw_text):
                global_findings.add(entity.raw_text)
            if self._normalized_contains(full_text, entity.raw_text):
                normalized_findings.add(entity.raw_text)

        return LeakageScanResult(
            leaked_values_detected=bool(global_findings or normalized_findings),
            findings=sorted(global_findings),
            findings_by_segment=segmented_findings,
            normalized_findings=sorted(normalized_findings),
        )

    @classmethod
    def _contains_surface(cls, text: str, surface: str) -> bool:
        pattern = cls._surface_pattern(surface)
        if pattern is None:
            return False
        return pattern.search(text) is not None

    @classmethod
    def _normalized_contains(cls, text: str, surface: str) -> bool:
        pattern = cls._normalized_surface_pattern(surface)
        if pattern is None:
            return False
        return pattern.search(cls._normalize_text(text)) is not None

    @classmethod
    def _surface_pattern(cls, surface: str) -> re.Pattern[str] | None:
        cleaned = " ".join(surface.split())
        if not cleaned:
            return None
        parts = cleaned.split(" ")
        escaped = [re.escape(part) for part in parts]
        return re.compile(
            rf"(?<![{cls._CONNECTOR_CHARS}]){'\\s+'.join(escaped)}(?![{cls._CONNECTOR_CHARS}])"
        )

    @classmethod
    def _normalized_surface_pattern(cls, surface: str) -> re.Pattern[str] | None:
        cleaned = cls._normalize_text(surface)
        if not cleaned:
            return None
        parts = cleaned.split(" ")
        escaped = [re.escape(part) for part in parts]
        return re.compile(rf"(?<!\w){'\\s+'.join(escaped)}(?!\w)")

    @staticmethod
    def _normalize_text(text: str) -> str:
        return re.sub(r"[^\w\s]", "", text.casefold())
